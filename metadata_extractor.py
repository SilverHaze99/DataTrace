#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Core Metadata Extraction Module
Extracts comprehensive metadata from various file types including:
- EXIF data from images
- Hidden strings and steganography detection
- File properties and system metadata
- Document metadata (PDF, Office, etc.)
- Audio metadata
- Archive analysis
"""
import os
import sys
import json
import struct
import hashlib
import mimetypes
import threading
import pickle
import shutil
from pathlib import Path
from datetime import datetime
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Optional, Union, Any
import io
import random
import zlib

# Required libraries
try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
try:
    import exifread
    EXIFREAD_AVAILABLE = True
except ImportError:
    EXIFREAD_AVAILABLE = False
try:
    from mutagen import File as MutagenFile
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
try:
    import zipfile
    import rarfile
    import tarfile
    ARCHIVE_AVAILABLE = True
except ImportError:
    ARCHIVE_AVAILABLE = False
    rarfile = None
try:
    import docx
    from openpyxl import load_workbook
    OFFICE_AVAILABLE = True
except ImportError:
    OFFICE_AVAILABLE = False
    docx = None
try:
    import subprocess
    EXIFTOOL_AVAILABLE = bool(shutil.which('exiftool'))
except ImportError:
    EXIFTOOL_AVAILABLE = False
try:
    import binascii
    BINASCII_AVAILABLE = True
except ImportError:
    BINASCII_AVAILABLE = False

from config import Config
from utils import calculate_entropy, human_readable_size, calculate_data_entropy, get_jpeg_marker_name, detect_languages

class MetadataExtractor:
    def __init__(self, config: Config = None, logger: logging.Logger = None):
        self.config = config or Config()
        self.logger = logger or logging.getLogger('MetadataExtractor')
        self.results = {}
        self.exiftool_cache = {}  # Cache for ExifTool results
        if not self.config.cache_dir.exists():
            self.config.cache_dir.mkdir(parents=True)

    def cached_exiftool(self, filepath: Path) -> Dict[str, Any]:
        """Centralized ExifTool execution with caching to reduce redundant calls."""
        cache_key = self.calculate_hashes(filepath).get('sha256', str(filepath))
        cache_file = self.config.cache_dir / f"{cache_key}.pickle"
        if cache_file.exists():
            try:
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                self.logger.warning(f"Cache read error for {filepath}: {e}")
        result = self.extract_with_exiftool(filepath)
        try:
            with open(cache_file, 'wb') as f:
                pickle.dump(result, f)
        except Exception as e:
            self.logger.warning(f"Cache write error for {filepath}: {e}")
        return result

    def extract_file_info(self, filepath: Path) -> Dict[str, Any]:
        """Extract basic file information with extended attributes."""
        try:
            stat = filepath.stat()
            extended_attrs = {}
            try:
                if hasattr(os, 'listxattr'):
                    attrs = os.listxattr(filepath)
                    for attr in attrs[:10]:
                        try:
                            value = os.getxattr(filepath, attr)
                            extended_attrs[attr] = value.decode('utf-8', errors='ignore')[:100]
                        except:
                            extended_attrs[attr] = "<binary_data>"
            except:
                pass
            file_info = {
                'filepath': str(filepath),
                'filename': filepath.name,
                'stem': filepath.stem,
                'suffix': filepath.suffix.lower(),
                'size_bytes': stat.st_size,
                'size_human': human_readable_size(stat.st_size),
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
                'mimetype': mimetypes.guess_type(filepath)[0],
                'permissions': oct(stat.st_mode)[-3:],
                'inode': getattr(stat, 'st_ino', None),
                'links': getattr(stat, 'st_nlink', None),
                'uid': getattr(stat, 'st_uid', None),
                'gid': getattr(stat, 'st_gid', None),
                'extended_attributes': extended_attrs if extended_attrs else None
            }
            if stat.st_size <= self.config.max_file_size_mb * 1024 * 1024:
                file_info.update(self.calculate_hashes(filepath))
            else:
                file_info['hash_skipped'] = f"File too large (>{self.config.max_file_size_mb}MB)"
            return file_info
        except Exception as e:
            self.logger.error(f"Error reading file info for {filepath}: {e}")
            return {'error': f"Error reading file info: {str(e)}"}

    def calculate_hashes(self, filepath: Path) -> Dict[str, str]:
        """Calculate MD5, SHA1, SHA256, and CRC32 hashes efficiently with verification."""
        try:
            md5_hash = hashlib.md5()
            sha1_hash = hashlib.sha1()
            sha256_hash = hashlib.sha256()
            crc32_hash = 0
            with open(filepath, 'rb') as f:
                while chunk := f.read(self.config.chunk_size):
                    md5_hash.update(chunk)
                    sha1_hash.update(chunk)
                    sha256_hash.update(chunk)
                    crc32_hash = zlib.crc32(chunk, crc32_hash)
            result = {
                'md5': md5_hash.hexdigest(),
                'sha1': sha1_hash.hexdigest(),
                'sha256': sha256_hash.hexdigest(),
                'crc32': format(crc32_hash, '08x'),
                'file_signature': sha256_hash.hexdigest()[:16]
            }
            if BINASCII_AVAILABLE:
                with open(filepath, 'rb') as f:
                    crc32_verify = 0
                    while chunk := f.read(self.config.chunk_size):
                        crc32_verify = binascii.crc32(chunk, crc32_verify)
                    if format(crc32_verify, '08x') != result['crc32']:
                        result['crc32_warning'] = "CRC32 verification mismatch"
            return result
        except Exception as e:
            self.logger.error(f"Error calculating hashes: {e}")
            return {'hash_error': str(e)}

    def extract_strings(self, filepath: Path, min_length: int = None,
                        max_strings: int = None) -> Dict[str, Any]:
        """Extract strings synchronously for memory efficiency."""
        min_length = min_length or self.config.min_string_length
        max_strings = max_strings or self.config.max_strings
        try:
            strings = []
            interesting_strings = []
            current_string = b''
            total_bytes = 0
            max_bytes = self.config.max_memory_mb * 1024 * 1024
            patterns = {
                'url': b'http',
                'email': b'@',
                'path': b'/',
                'registry': b'HKEY',
                'base64': b'==',
            }
            with open(filepath, 'rb') as f:
                while total_bytes < max_bytes:
                    chunk = f.read(self.config.chunk_size)
                    if not chunk:
                        break
                    for byte in chunk:
                        total_bytes += 1
                        if 32 <= byte <= 126 or byte >= 128:
                            current_string += bytes([byte])
                        else:
                            if len(current_string) >= min_length:
                                try:
                                    decoded = current_string.decode('utf-8', errors='ignore').strip()
                                    if decoded:
                                        strings.append(decoded)
                                        decoded_lower = decoded.lower()
                                        for pattern_name, pattern in patterns.items():
                                            if pattern.decode('utf-8', errors='ignore') in decoded_lower:
                                                interesting_strings.append({
                                                    'type': pattern_name,
                                                    'string': decoded[:200],
                                                    'offset': total_bytes - len(current_string)
                                                })
                                        if len(strings) >= max_strings:
                                            break
                                except UnicodeDecodeError:
                                    pass
                            current_string = b''
                            if len(strings) >= max_strings:
                                break
                        if len(strings) >= max_strings:
                            break
                    if len(strings) >= max_strings:
                        break
            if current_string and len(current_string) >= min_length:
                try:
                    decoded = current_string.decode('utf-8', errors='ignore').strip()
                    if decoded:
                        strings.append(decoded)
                except UnicodeDecodeError:
                    pass
            return {
                'total_strings': len(strings),
                'strings': strings[:100],
                'interesting_strings': interesting_strings[:50],
                'bytes_analyzed': total_bytes,
                'truncated': total_bytes >= max_bytes
            }
        except Exception as e:
            self.logger.error(f"Error extracting strings: {e}")
            return {'strings_error': str(e)}

    def detect_steganography(self, filepath: Path) -> Dict[str, Any]:
        """Detect steganography through size analysis, embedded signatures, and LSB analysis."""
        if not self.config.enable_steganography:
            return {}
        try:
            results = {}
            file_size = filepath.stat().st_size
            if filepath.suffix.lower() in ['.jpg', '.jpeg', '.png', '.gif'] and PIL_AVAILABLE:
                try:
                    with Image.open(filepath) as img:
                        expected_size = img.width * img.height * 3
                        ratio = file_size / expected_size if expected_size > 0 else 0
                        results['size_analysis'] = {
                            'actual_size': file_size,
                            'expected_size': expected_size,
                            'ratio': ratio,
                            'suspicious': ratio > 1.5
                        }
                        if img.mode == 'RGB':
                            pixels = img.convert('RGB').getdata()
                            lsb_counts = [0] * 2
                            sample_size = min(len(pixels), 10000)
                            for pixel in random.sample(list(pixels), sample_size):
                                for channel in pixel:
                                    lsb_counts[channel & 1] += 1
                            lsb_ratio = lsb_counts[1] / sum(lsb_counts) if sum(lsb_counts) > 0 else 0
                            results['lsb_analysis'] = {
                                'lsb_ratio': lsb_ratio,
                                'suspicious': abs(lsb_ratio - 0.5) > 0.1
                            }
                except Exception as e:
                    results['image_error'] = str(e)
            with open(filepath, 'rb') as f:
                data = f.read(min(file_size, 1024 * 1024))
                signatures = {
                    b'PK\x03\x04': 'ZIP archive',
                    b'\xff\xd8\xff': 'JPEG image',
                    b'\x89PNG\r\n\x1a\n': 'PNG image',
                    b'GIF8': 'GIF image',
                    b'%PDF': 'PDF document',
                    b'RIFF': 'RIFF container',
                }
                found_signatures = []
                for sig, desc in signatures.items():
                    pos = data.find(sig, 100)
                    if pos > 0:
                        found_signatures.append({
                            'signature': sig.hex(),
                            'description': desc,
                            'offset': pos
                        })
                if found_signatures:
                    results['embedded_signatures'] = found_signatures
            if file_size < 10 * 1024 * 1024:
                entropy = calculate_entropy(filepath)
                results['entropy'] = {
                    'value': entropy,
                    'suspicious': entropy > 7.5
                }
            return results
        except Exception as e:
            self.logger.error(f"Error in steganography detection: {e}")
            return {'stego_error': str(e)}

    def extract_image_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract comprehensive image metadata using multiple libraries."""
        metadata = {}
        if PIL_AVAILABLE:
            try:
                with Image.open(filepath) as img:
                    metadata['image_info'] = {
                        'format': img.format,
                        'mode': img.mode,
                        'size': img.size,
                        'width': img.width,
                        'height': img.height,
                        'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info,
                        'animation': getattr(img, 'is_animated', False),
                        'frames': getattr(img, 'n_frames', 1)
                    }
                    exif_data = getattr(img, '_getexif', lambda: None)()
                    if exif_data:
                        exif_readable = {}
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            if isinstance(value, (str, int, float)):
                                exif_readable[tag] = value
                            elif isinstance(value, bytes):
                                try:
                                    exif_readable[tag] = value.decode('utf-8', errors='ignore')
                                except:
                                    exif_readable[tag] = value.hex()[:100]
                            else:
                                exif_readable[tag] = str(value)[:200]
                        metadata['exif_pil'] = exif_readable
                    try:
                        colors = img.getcolors(maxcolors=256)
                        if colors:
                            metadata['color_analysis'] = {
                                'unique_colors': len(colors),
                                'most_common_color': colors[0][1] if colors else None,
                                'is_grayscale': img.mode in ('L', 'LA', '1')
                            }
                    except:
                        pass
                    if 'icc_profile' in img.info:
                        metadata['icc_profile_present'] = True
                        metadata['icc_profile_size'] = len(img.info['icc_profile'])
                        icc_data = img.info['icc_profile']
                        if len(icc_data) >= 128:
                            metadata['icc_profile_info'] = {
                                'size': len(icc_data),
                                'cmm_type': icc_data[4:8].decode('ascii', errors='ignore').strip(),
                                'version': f"{icc_data[8]}.{icc_data[9]}.{icc_data[10]}",
                                'device_class': icc_data[12:16].decode('ascii', errors='ignore').strip(),
                                'color_space': icc_data[16:20].decode('ascii', errors='ignore').strip(),
                            }
            except Exception as e:
                self.logger.error(f"PIL image analysis error: {e}")
                metadata['pil_error'] = str(e)
        if EXIFREAD_AVAILABLE:
            try:
                with open(filepath, 'rb') as f:
                    tags = exifread.process_file(f, details=True)
                    exif_detailed = {}
                    for tag in tags.keys():
                        if tag not in ['JPEGThumbnail', 'TIFFThumbnail', 'Filename']:
                            try:
                                value = str(tags[tag])
                                if len(value) < 500:
                                    exif_detailed[tag] = value
                            except:
                                exif_detailed[tag] = f"<{type(tags[tag]).__name__}>"
                    if exif_detailed:
                        metadata['exif_detailed'] = exif_detailed
            except Exception as e:
                self.logger.error(f"exifread error: {e}")
                metadata['exifread_error'] = str(e)
        elif EXIFTOOL_AVAILABLE:
            metadata.update(self.cached_exiftool(filepath))
        if filepath.suffix.lower() in ['.jpg', '.jpeg']:
            metadata.update(self.extract_jpeg_segments(filepath))
        return metadata

    def extract_with_exiftool(self, filepath: Path) -> Dict[str, Any]:
        """Use ExifTool for comprehensive metadata extraction."""
        if not EXIFTOOL_AVAILABLE:
            return {'exiftool_info': 'ExifTool not installed (recommended for complete metadata)'}
        try:
            cmd = ['exiftool', '-j', '-n', '-struct', '-duplicates', str(filepath)]
            result = subprocess.run(cmd, capture_output=True, text=True,
                                   timeout=self.config.timeout_seconds)
            if result.returncode == 0:
                try:
                    exiftool_data = json.loads(result.stdout)
                    if exiftool_data and len(exiftool_data) > 0:
                        return {'exiftool_metadata': exiftool_data[0]}
                except json.JSONDecodeError as e:
                    return {'exiftool_error': f'JSON parse error: {e}'}
            else:
                return {'exiftool_error': f'ExifTool error: {result.stderr}'}
        except subprocess.TimeoutExpired:
            return {'exiftool_error': 'ExifTool timeout'}
        except Exception as e:
            return {'exiftool_error': str(e)}

    def extract_jpeg_segments(self, filepath: Path) -> Dict[str, Any]:
        """Analyze JPEG segments for hidden data and structure."""
        try:
            segments = []
            with open(filepath, 'rb') as f:
                marker = f.read(2)
                if marker != b'\xff\xd8':
                    return {'jpeg_error': 'Not a valid JPEG file'}
                segments.append({
                    'marker': 'FFD8',
                    'type': get_jpeg_marker_name(0xd8),
                    'size': 0,
                    'offset': 0
                })
                while True:
                    marker = f.read(2)
                    if len(marker) != 2:
                        break
                    if marker[0] != 0xff:
                        break
                    marker_type = marker[1]
                    current_offset = f.tell() - 2
                    if marker_type in [0xd0, 0xd1, 0xd2, 0xd3, 0xd4, 0xd5, 0xd6, 0xd7, 0xd8, 0xd9]:
                        segments.append({
                            'marker': f'FF{marker_type:02X}',
                            'type': get_jpeg_marker_name(marker_type),
                            'size': 0,
                            'offset': current_offset
                        })
                        continue
                    length_bytes = f.read(2)
                    if len(length_bytes) != 2:
                        break
                    length = struct.unpack('>H', length_bytes)[0] - 2
                    segment_data = f.read(length) if length > 0 else b''
                    segment_info = {
                        'marker': f'FF{marker_type:02X}',
                        'type': get_jpeg_marker_name(marker_type),
                        'size': length,
                        'offset': current_offset
                    }
                    if marker_type == 0xe1 and segment_data.startswith(b'Exif\x00\x00'):
                        segment_info['content'] = 'EXIF data'
                        segment_info['exif_size'] = len(segment_data) - 6
                    elif marker_type == 0xe2 and segment_data.startswith(b'ICC_PROFILE'):
                        segment_info['content'] = 'ICC Profile'
                        segment_info['icc_size'] = len(segment_data)
                    elif marker_type == 0xe0 and segment_data.startswith(b'JFIF'):
                        segment_info['content'] = 'JFIF Header'
                        if len(segment_data) >= 9:
                            version = f"{segment_data[5]}.{segment_data[6]:02d}"
                            segment_info['jfif_version'] = version
                    elif marker_type == 0xed and b'Photoshop' in segment_data:
                        segment_info['content'] = 'Photoshop Data'
                        segment_info['ps_size'] = len(segment_data)
                    elif marker_type == 0xfe:
                        try:
                            comment = segment_data.decode('utf-8', errors='ignore').strip()
                            segment_info['content'] = f'Comment: {comment[:100]}'
                        except:
                            segment_info['content'] = 'Comment (binary data)'
                    elif length > 1000:
                        segment_info['large_segment'] = True
                        segment_info['entropy'] = calculate_data_entropy(segment_data[:1000])
                    segments.append(segment_info)
                    if marker_type == 0xda:
                        break
            return {
                'jpeg_segments': segments,
                'total_segments': len(segments),
                'suspicious_segments': [s for s in segments if s.get('large_segment') or s.get('entropy', 0) > 7.0]
            }
        except Exception as e:
            self.logger.error(f"JPEG segment analysis error: {e}")
            return {'jpeg_segments_error': str(e)}

    def extract_audio_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract comprehensive audio metadata."""
        if not MUTAGEN_AVAILABLE:
            if EXIFTOOL_AVAILABLE:
                return self.cached_exiftool(filepath)
            return {'error': 'Mutagen library not available'}
        try:
            audio_file = MutagenFile(str(filepath))
            if audio_file is None:
                return {'error': 'No audio metadata found or unsupported format'}
            metadata = {}
            if hasattr(audio_file, 'info'):
                info = audio_file.info
                metadata['audio_info'] = {
                    'length_seconds': getattr(info, 'length', None),
                    'bitrate': getattr(info, 'bitrate', None),
                    'sample_rate': getattr(info, 'sample_rate', None),
                    'channels': getattr(info, 'channels', None),
                    'format': audio_file.mime[0] if audio_file.mime else None,
                    'codec': getattr(info, 'codec', None),
                    'bitrate_mode': getattr(info, 'bitrate_mode', None),
                }
                if info.length and info.bitrate:
                    estimated_size = (info.length * info.bitrate * 1000) // 8
                    actual_size = filepath.stat().st_size
                    metadata['audio_info']['size_efficiency'] = actual_size / estimated_size if estimated_size > 0 else 0
            if audio_file.tags:
                tags = {}
                for key, value in audio_file.tags.items():
                    if isinstance(value, list):
                        tags[key] = [str(v)[:200] for v in value]
                    else:
                        tags[key] = str(value)[:200]
                metadata['tags'] = tags
                common_tags = {
                    'title': ['TIT2', 'TITLE', '©nam'],
                    'artist': ['TPE1', 'ARTIST', '©ART'],
                    'album': ['TALB', 'ALBUM', '©alb'],
                    'date': ['TDRC', 'DATE', '©day'],
                    'genre': ['TCON', 'GENRE', '©gen'],
                }
                normalized_tags = {}
                for tag_name, possible_keys in common_tags.items():
                    for key in possible_keys:
                        if key in tags:
                            normalized_tags[tag_name] = tags[key]
                            break
                if normalized_tags:
                    metadata['common_tags'] = normalized_tags
            if EXIFTOOL_AVAILABLE:
                metadata.update(self.cached_exiftool(filepath))
            return metadata
        except Exception as e:
            self.logger.error(f"Audio metadata extraction error: {e}")
            return {'audio_error': str(e)}

    def extract_pdf_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract comprehensive PDF metadata."""
        if not PYPDF2_AVAILABLE:
            if EXIFTOOL_AVAILABLE:
                return self.cached_exiftool(filepath)
            return {'error': 'PyPDF2 library not available'}
        try:
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted
                }
                if pdf_reader.metadata:
                    pdf_meta = {}
                    for key, value in pdf_reader.metadata.items():
                        clean_key = key.replace('/', '') if key.startswith('/') else key
                        pdf_meta[clean_key] = str(value)[:500]
                    metadata['document_info'] = pdf_meta
                try:
                    text_sample = ""
                    for i, page in enumerate(pdf_reader.pages[:3]):
                        page_text = page.extract_text()
                        text_sample += page_text[:1000]
                        if len(text_sample) > 2000:
                            break
                    if text_sample.strip():
                        metadata['text_analysis'] = {
                            'has_extractable_text': True,
                            'sample_length': len(text_sample),
                            'languages_detected': detect_languages(text_sample),
                            'word_count_estimate': len(text_sample.split())
                        }
                    else:
                        metadata['text_analysis'] = {'has_extractable_text': False}
                except:
                    metadata['text_analysis'] = {'extraction_error': 'Could not extract text'}
                try:
                    form_fields = []
                    annotations = []
                    for page_num, page in enumerate(pdf_reader.pages[:10]):
                        if '/Annots' in page:
                            annotations.append(f'Page {page_num + 1} has annotations')
                        if '/AcroForm' in pdf_reader.trailer.get('/Root', {}):
                            form_fields.append('Document contains form fields')
                    if form_fields or annotations:
                        metadata['interactive_elements'] = {
                            'form_fields': form_fields,
                            'annotations': annotations
                        }
                except:
                    pass
                if EXIFTOOL_AVAILABLE:
                    metadata.update(self.cached_exiftool(filepath))
                return metadata
        except Exception as e:
            self.logger.error(f"PDF metadata extraction error: {e}")
            return {'pdf_error': str(e)}

    def extract_office_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract metadata from Office documents."""
        if not OFFICE_AVAILABLE:
            if EXIFTOOL_AVAILABLE:
                return self.cached_exiftool(filepath)
            return {'error': 'Office libraries not available (python-docx, openpyxl)'}
        try:
            metadata = {}
            suffix = filepath.suffix.lower()
            if suffix == '.docx' and docx:
                try:
                    doc = docx.Document(str(filepath))
                    core_props = doc.core_properties
                    metadata['core_properties'] = {
                        'title': core_props.title,
                        'author': core_props.author,
                        'subject': core_props.subject,
                        'keywords': core_props.keywords,
                        'comments': core_props.comments,
                        'created': core_props.created.isoformat() if core_props.created else None,
                        'modified': core_props.modified.isoformat() if core_props.modified else None,
                        'last_modified_by': core_props.last_modified_by,
                        'category': core_props.category,
                        'language': core_props.language,
                        'revision': core_props.revision,
                    }
                    metadata['document_stats'] = {
                        'paragraphs': len(doc.paragraphs),
                        'sections': len(doc.sections),
                        'tables': len(doc.tables),
                    }
                    full_text = []
                    for paragraph in doc.paragraphs[:20]:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                    text_content = ' '.join(full_text)
                    if text_content:
                        metadata['text_analysis'] = {
                            'word_count': len(text_content.split()),
                            'char_count': len(text_content),
                            'languages_detected': detect_languages(text_content[:1000])
                        }
                except Exception as e:
                    metadata['docx_error'] = str(e)
            elif suffix in ['.xlsx', '.xls']:
                try:
                    workbook = load_workbook(str(filepath), read_only=True, data_only=True)
                    props = workbook.properties
                    metadata['workbook_properties'] = {
                        'title': props.title,
                        'creator': props.creator,
                        'subject': props.subject,
                        'description': props.description,
                        'keywords': props.keywords,
                        'created': props.created.isoformat() if props.created else None,
                        'modified': props.modified.isoformat() if props.modified else None,
                        'last_modified_by': props.lastModifiedBy,
                        'category': props.category,
                        'language': props.language,
                        'version': props.version,
                    }
                    sheets_info = []
                    for sheet_name in workbook.sheetnames[:10]:
                        sheet = workbook[sheet_name]
                        sheet_info = {
                            'name': sheet_name,
                            'max_row': sheet.max_row,
                            'max_column': sheet.max_column,
                            'has_data': sheet.max_row > 1 or sheet.max_column > 1
                        }
                        formula_count = 0
                        cell_count = 0
                        for row in sheet.iter_rows(max_row=min(10, sheet.max_row),
                                                  max_col=min(10, sheet.max_column)):
                            for cell in row:
                                if cell.data_type == 'f':
                                    formula_count += 1
                                cell_count += 1
                                if cell_count >= 100:
                                    break
                            if cell_count >= 100:
                                break
                        sheet_info['formula_ratio'] = formula_count / max(cell_count, 1)
                        sheets_info.append(sheet_info)
                    metadata['sheets_analysis'] = {
                        'total_sheets': len(workbook.sheetnames),
                        'sheets_info': sheets_info
                    }
                    workbook.close()
                except Exception as e:
                    metadata['xlsx_error'] = str(e)
            if EXIFTOOL_AVAILABLE:
                metadata.update(self.cached_exiftool(filepath))
            return metadata
        except Exception as e:
            self.logger.error(f"Office metadata extraction error: {e}")
            return {'office_error': str(e)}

    def extract_archive_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract metadata from archive files with zip-bomb protection."""
        try:
            metadata = {}
            suffix = filepath.suffix.lower()
            if suffix == '.zip':
                try:
                    with zipfile.ZipFile(filepath, 'r') as zip_file:
                        file_list = zip_file.filelist
                        if len(file_list) > 10000 or self.check_zip_bomb(zip_file):
                            metadata['security'] = {'zip_bomb_warning': 'Potential zip bomb detected'}
                            return metadata
                        metadata['archive_info'] = {
                            'type': 'ZIP',
                            'total_files': len(file_list),
                            'compressed_size': sum(f.compress_size for f in file_list),
                            'uncompressed_size': sum(f.file_size for f in file_list),
                        }
                        if metadata['archive_info']['compressed_size'] > 0:
                            compression_ratio = metadata['archive_info']['uncompressed_size'] / metadata['archive_info']['compressed_size']
                            metadata['archive_info']['compression_ratio'] = compression_ratio
                        file_types = {}
                        suspicious_files = []
                        for file_info in file_list[:100]:
                            if file_info.filename.endswith('/'):
                                continue
                            ext = Path(file_info.filename).suffix.lower()
                            file_types[ext] = file_types.get(ext, 0) + 1
                            if any(pattern in file_info.filename.lower() for pattern in
                                  ['password', 'crack', 'keygen', 'patch', '.exe', '.bat', '.cmd']):
                                suspicious_files.append(file_info.filename)
                            if file_info.file_size > 0:
                                file_ratio = file_info.compress_size / file_info.file_size
                                if file_ratio > 0.95 and file_info.file_size > 1000:
                                    suspicious_files.append(f"{file_info.filename} (poor compression)")
                        metadata['content_analysis'] = {
                            'file_types': file_types,
                            'suspicious_files': suspicious_files[:10]
                        }
                        encrypted_files = [f.filename for f in file_list if f.flag_bits & 0x1]
                        if encrypted_files:
                            metadata['security'] = {
                                'encrypted_files': len(encrypted_files),
                                'encrypted_file_names': encrypted_files[:10]
                            }
                except Exception as e:
                    metadata['zip_error'] = str(e)
            elif suffix == '.rar' and rarfile:
                try:
                    with rarfile.RarFile(filepath) as rar_file:
                        file_list = rar_file.infolist()
                        if len(file_list) > 10000:
                            metadata['security'] = {'rar_bomb_warning': 'Potential archive bomb detected'}
                            return metadata
                        metadata['archive_info'] = {
                            'type': 'RAR',
                            'total_files': len(file_list),
                            'compressed_size': sum(f.compress_size for f in file_list),
                            'uncompressed_size': sum(f.file_size for f in file_list),
                        }
                        file_types = {}
                        for file_info in file_list[:100]:
                            if not file_info.is_dir():
                                ext = Path(file_info.filename).suffix.lower()
                                file_types[ext] = file_types.get(ext, 0) + 1
                        metadata['content_analysis'] = {'file_types': file_types}
                except Exception as e:
                    metadata['rar_error'] = str(e)
            elif suffix == '.gz' and tarfile:
                try:
                    with tarfile.open(filepath, 'r:gz') as tar_file:
                        file_list = tar_file.getmembers()
                        if len(file_list) > 10000:
                            metadata['security'] = {'tar_bomb_warning': 'Potential archive bomb detected'}
                            return metadata
                        metadata['archive_info'] = {
                            'type': 'TAR.GZ',
                            'total_files': len(file_list),
                            'uncompressed_size': sum(f.size for f in file_list if not f.isdir()),
                        }
                        file_types = {}
                        for file_info in file_list[:100]:
                            if not file_info.isdir():
                                ext = Path(file_info.name).suffix.lower()
                                file_types[ext] = file_types.get(ext, 0) + 1
                        metadata['content_analysis'] = {'file_types': file_types}
                except Exception as e:
                    metadata['tar_error'] = str(e)
            else:
                self.logger.warning(f"RAR/TAR.GZ support limited without rarfile/tarfile")
                if EXIFTOOL_AVAILABLE:
                    metadata.update(self.cached_exiftool(filepath))
            if EXIFTOOL_AVAILABLE and not metadata.get('security'):
                metadata.update(self.cached_exiftool(filepath))
            return metadata
        except Exception as e:
            self.logger.error(f"Archive metadata extraction error: {e}")
            return {'archive_error': str(e)}

    def check_zip_bomb(self, zip_file: zipfile.ZipFile) -> bool:
        """Check for potential zip bombs."""
        total_size = sum(f.file_size for f in zip_file.filelist)
        if total_size > 1024 * 1024 * 1024 * 10:  # 10GB uncompressed
            return True
        nesting_level = 0
        for file_info in zip_file.filelist:
            if file_info.filename.endswith('.zip'):
                nesting_level += 1
                if nesting_level > 5:
                    return True
        return False

    def extract_hex_header(self, filepath: Path, bytes_count: int = 64) -> Dict[str, Any]:
        """Extract and analyze hex header with enhanced analysis."""
        try:
            with open(filepath, 'rb') as f:
                header_bytes = f.read(bytes_count)
                hex_string = ' '.join(f'{b:02x}' for b in header_bytes)
                ascii_string = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in header_bytes)
                magic_signatures = {
                    b'\xff\xd8\xff': 'JPEG image',
                    b'\x89PNG\r\n\x1a\n': 'PNG image',
                    b'GIF8': 'GIF image',
                    b'RIFF': 'RIFF container (AVI, WAV)',
                    b'%PDF': 'PDF document',
                    b'PK\x03\x04': 'ZIP archive',
                    b'PK\x05\x06': 'ZIP archive (empty)',
                    b'PK\x07\x08': 'ZIP archive (spanned)',
                    b'Rar!\x1a\x07\x00': 'RAR archive',
                    b'\x7fELF': 'ELF executable',
                    b'MZ': 'PE executable',
                    b'\xca\xfe\xba\xbe': 'Java class file',
                    b'FLAC': 'FLAC audio',
                    b'ID3': 'MP3 audio',
                    b'\xff\xfb': 'MP3 audio',
                    b'OggS': 'OGG container',
                    b'ftyp': 'MP4 container (offset 4)',
                    b'\x00\x00\x00\x18ftypmp4': 'MP4 video',
                    b'\x1f\x8b\x08': 'GZIP compressed',
                    b'BZh': 'BZIP2 compressed',
                    b'\xfd7zXZ\x00': 'XZ compressed',
                }
                detected_type = None
                for signature, description in magic_signatures.items():
                    if header_bytes.startswith(signature):
                        detected_type = description
                        break
                    elif len(header_bytes) > 4 and header_bytes[4:].startswith(signature):
                        detected_type = description + " (offset 4)"
                        break
                return {
                    'hex': hex_string,
                    'ascii': ascii_string,
                    'magic_bytes': header_bytes[:8].hex(),
                    'detected_type': detected_type,
                    'header_entropy': calculate_data_entropy(header_bytes),
                    'null_bytes': header_bytes.count(0),
                    'printable_ratio': sum(1 for b in header_bytes if 32 <= b <= 126) / len(header_bytes)
                }
        except Exception as e:
            self.logger.error(f"Hex header analysis error: {e}")
            return {'hex_error': str(e)}

    def analyze_file(self, filepath: Union[str, Path], extract_strings: bool = True,
                    string_min_length: int = None, filter_types: List[str] = None) -> Dict[str, Any]:
        """Main function for comprehensive file analysis with optional filtering."""
        filepath = Path(filepath)
        if not filepath.exists():
            return {'error': f'File not found: {filepath}'}
        if not filepath.is_file():
            return {'error': f'Path is not a file: {filepath}'}
        self.logger.info(f"Analyzing file: {filepath}")
        results = {'analysis_timestamp': datetime.now().isoformat()}
        try:
            if not filter_types or 'file_info' in filter_types:
                results['file_info'] = self.extract_file_info(filepath)
            if not filter_types or 'hex_header' in filter_types:
                results['hex_header'] = self.extract_hex_header(filepath)
            if extract_strings and (not filter_types or 'strings' in filter_types):
                results['strings'] = self.extract_strings(
                    filepath, string_min_length or self.config.min_string_length
                )
            if self.config.enable_steganography and (not filter_types or 'steganography' in filter_types):
                results['steganography'] = self.detect_steganography(filepath)
            mime_type = results['file_info'].get('mimetype', '') if 'file_info' in results else ''
            extension = results['file_info'].get('suffix', '').lower() if 'file_info' in results else filepath.suffix.lower()
            if mime_type and mime_type.startswith('image/') and (not filter_types or 'image_metadata' in filter_types):
                self.logger.info("Performing image metadata extraction")
                results['image_metadata'] = self.extract_image_metadata(filepath)
            elif mime_type and mime_type.startswith('audio/') and (not filter_types or 'audio_metadata' in filter_types):
                self.logger.info("Performing audio metadata extraction")
                results['audio_metadata'] = self.extract_audio_metadata(filepath)
            elif (extension == '.pdf' or (mime_type and 'pdf' in mime_type)) and (not filter_types or 'pdf_metadata' in filter_types):
                self.logger.info("Performing PDF metadata extraction")
                results['pdf_metadata'] = self.extract_pdf_metadata(filepath)
            elif extension in ['.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt'] and (not filter_types or 'office_metadata' in filter_types):
                self.logger.info("Performing Office document metadata extraction")
                results['office_metadata'] = self.extract_office_metadata(filepath)
            elif extension in ['.zip', '.rar', '.7z', '.tar', '.gz'] and (not filter_types or 'archive_metadata' in filter_types):
                self.logger.info("Performing archive metadata extraction")
                results['archive_metadata'] = self.extract_archive_metadata(filepath)
            else:
                results['generic_analysis'] = {
                    'file_type': 'Unknown/Binary',
                    'requires_specialized_tools': True
                }
            self.logger.info(f"Analysis completed for: {filepath}")
            return results
        except Exception as e:
            self.logger.error(f"Critical error during analysis: {e}")
            return {'critical_error': str(e), 'partial_results': results}

    def analyze_multiple_files(self, filepaths: List[Union[str, Path]],
                             progress_callback=None, filter_types: List[str] = None) -> Dict[str, Dict[str, Any]]:
        """Analyze multiple files with threading support."""
        results = {}
        with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
            future_to_path = {
                executor.submit(self.analyze_file, path, filter_types=filter_types): path
                for path in filepaths
            }
            for future in as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    result = future.result(timeout=self.config.timeout_seconds * 2)
                    results[str(path)] = result
                except Exception as e:
                    self.logger.error(f"Error analyzing {path}: {e}")
                    results[str(path)] = {'error': str(e)}
                if progress_callback:
                    progress_callback()
        return results